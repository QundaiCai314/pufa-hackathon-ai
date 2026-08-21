from collections import Counter
from io import BytesIO
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document
from sqlalchemy import text
from app.api.auth import db, current_user

router = APIRouter(tags=["admin"])

def admin_user(user=Depends(current_user)):
    if not user.get("is_admin") and user.get("user_type") != "admin":
        raise HTTPException(403, "仅管理用户可访问")
    return user

@router.get("/overview")
def overview(user=Depends(admin_user)):
    with db().connect() as conn:
        users = conn.execute(text("SELECT COUNT(*) FROM users WHERE is_active=TRUE")).scalar() or 0
        sessions = conn.execute(text("SELECT COUNT(*) FROM chat_sessions")).scalar() or 0
        messages = conn.execute(text("SELECT COUNT(*) FROM chat_messages")).scalar() or 0
        levels = conn.execute(text("SELECT lead_level, COUNT(*) AS count FROM chat_sessions GROUP BY lead_level")).mappings().all()
        roles = conn.execute(text("SELECT assistant_role, COUNT(*) AS count FROM chat_sessions GROUP BY assistant_role")).mappings().all()
        recent = conn.execute(text("SELECT s.id,s.session_name,s.lead_score,s.lead_level,s.updated_at,u.username FROM chat_sessions s JOIN users u ON u.id=s.user_id ORDER BY s.updated_at DESC LIMIT 20")).mappings().all()
    return {"users": users, "sessions": sessions, "messages": messages, "lead_distribution": [dict(x) for x in levels], "role_distribution": [dict(x) for x in roles], "recent_sessions": [dict(x) for x in recent]}

@router.get("/project-health")
def project_health(user=Depends(admin_user)):
    """管理员项目健康度：仅依据会话、已保存方案和可追溯资料计算。"""
    required = ["应用场景", "目标规模", "压力要求", "部署方式", "能源来源"]
    with db().connect() as conn:
        sessions = conn.execute(text("""SELECT s.id,s.session_name,s.lead_score,s.updated_at,u.username,
            p.payload FROM chat_sessions s JOIN users u ON u.id=s.user_id
            LEFT JOIN LATERAL (SELECT payload FROM proposal_versions WHERE session_id=s.id ORDER BY version_no DESC LIMIT 1) p ON TRUE
            ORDER BY s.updated_at DESC LIMIT 20""")).mappings().all()
    items = []
    for session in sessions:
        payload = session["payload"] or {}
        if isinstance(payload, str):
            import json; payload = json.loads(payload)
        profile = payload.get("profile") or {}
        evidence = payload.get("normalized_rows") or []
        conflicts = payload.get("conflicts") or []
        completeness = round(100 * sum(bool(profile.get(k)) for k in required) / len(required))
        evidence_score = min(100, len(evidence) * 20) if evidence else 25
        field_score = 100 if profile.get("部署方式") else 35
        consistency = max(0, 100 - len(conflicts) * 35)
        match = min(100, max(25, int(session["lead_score"] or 0) + (20 if profile.get("目标规模") else 0) + (15 if evidence else 0)))
        feasibility = round((completeness + field_score + consistency) / 3)
        health = round(completeness*.22 + match*.20 + evidence_score*.20 + field_score*.13 + feasibility*.15 + consistency*.10)
        risks = []
        if not profile.get("压力要求"): risks.append({"level":"高", "text":"出口/用氢压力未确认"})
        if not profile.get("部署方式"): risks.append({"level":"中", "text":"现场部署方式未明确"})
        if conflicts: risks.append({"level":"中", "text":f"{len(conflicts)} 项参数存在资料冲突，待人工核验"})
        if evidence: risks.append({"level":"低", "text":"产品参数已保留可追溯资料来源"})
        items.append({"session_id":str(session["id"]),"project":session["session_name"] or "未命名项目","customer":session["username"],"updated_at":session["updated_at"],"health":health,"dimensions":{"需求完整度":completeness,"产品匹配度":match,"参数可核验度":evidence_score,"现场条件明确度":field_score,"部署可行性":feasibility,"资料一致性":consistency},"risks":risks,"has_proposal":bool(payload)})
    return {"projects": items}

@router.get("/users")
def users(user=Depends(admin_user)):
    with db().connect() as conn:
        rows=conn.execute(text("SELECT id,username,email,user_type,is_admin,is_active,created_at FROM users ORDER BY created_at DESC")).mappings().all()
    return {"users":[dict(x) for x in rows]}

@router.get("/leads")
def leads(user=Depends(admin_user)):
    with db().connect() as conn:
        rows=conn.execute(text("SELECT s.id,s.session_name,s.lead_score,s.lead_level,s.lead_signals,s.updated_at,u.username,u.email FROM chat_sessions s JOIN users u ON u.id=s.user_id ORDER BY s.lead_score DESC,s.updated_at DESC")).mappings().all()
    return {"leads":[dict(x) for x in rows]}

@router.get("/sessions/{session_id}")
def session_detail(session_id: str, user=Depends(admin_user)):
    with db().connect() as conn:
        rows=conn.execute(text("SELECT m.id,m.role,m.content,m.created_at,u.username FROM chat_messages m JOIN chat_sessions s ON s.id=m.session_id JOIN users u ON u.id=s.user_id WHERE s.id=:s ORDER BY m.created_at"), {"s":session_id}).mappings().all()
    return {"messages":[dict(x) for x in rows]}

@router.get("/leads/{session_id}/sales-plan.docx")
def sales_plan(session_id: str, user=Depends(admin_user)):
    with db().connect() as conn:
        lead = conn.execute(text("SELECT s.session_name,s.lead_score,s.lead_level,u.username,u.email FROM chat_sessions s JOIN users u ON u.id=s.user_id WHERE s.id=:s"), {"s":session_id}).mappings().first()
        rows = conn.execute(text("SELECT role,content,created_at FROM chat_messages WHERE session_id=:s ORDER BY created_at"), {"s":session_id}).mappings().all()
    if not lead: raise HTTPException(404, "线索不存在")
    doc = Document()
    doc.add_heading('客户销售沟通方案', 0)
    doc.add_paragraph(f"客户：{lead['username']}  |  邮箱：{lead['email'] or '未提供'}")
    doc.add_paragraph(f"线索等级：{lead['lead_level']}  |  意向分：{lead['lead_score']}")
    doc.add_heading('沟通话术建议', level=1)
    doc.add_paragraph('您好，感谢您关注氢璞解决方案。基于您在沟通中提到的项目需求，我们建议先确认应用场景、目标规模、关键技术约束与项目时间节点，再由销售与技术团队提供匹配的产品方案和参数说明。')
    doc.add_heading('建议下一步', level=1)
    for item in ['确认项目应用场景与目标产能/功率', '确认压力、纯度、能源来源及部署方式', '安排技术交流并输出正式选型方案', '根据确认后的配置提供商务报价与交付计划']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_heading('客户对话依据', level=1)
    for r in rows[-20:]:
        doc.add_paragraph(f"{'客户' if r['role']=='user' else 'AI'}：{r['content']}")
    out=BytesIO(); doc.save(out); out.seek(0)
    name=f"销售方案_{lead['username']}.docx"
    return StreamingResponse(out, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={'Content-Disposition': f"attachment; filename*=UTF-8''{name}"})
