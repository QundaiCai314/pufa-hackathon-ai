import os
import json
import os
from datetime import datetime, timedelta, timezone
from typing import Optional
from fastapi import APIRouter, HTTPException, Depends, Header
from fastapi.responses import StreamingResponse
from io import BytesIO
from docx import Document
import zipfile
import html
import fitz
from pydantic import BaseModel
from jose import jwt, JWTError
from passlib.context import CryptContext
from sqlalchemy import create_engine, text

router = APIRouter(tags=["auth"])
DATABASE_URL = os.getenv("DATABASE_URL", "")
SECRET_KEY = os.getenv("JWT_SECRET_KEY", "qingpu-change-this-secret")
ALGORITHM = "HS256"
pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")

class ProposalSaveRequest(BaseModel):
    title: str = "项目技术方案"
    profile: dict = {}
    content: str = ""
    results: list = []

STANDARD_FIELDS = {
    "额定产氢量": "额定产氢量", "产氢量": "额定产氢量", "输出功率": "输出功率",
    "峰值功率": "峰值输出功率", "工作压力": "工作压力", "氢气产品压力": "氢气产品压力",
    "运行温度": "工作温度", "工作温度": "工作温度", "设计寿命": "设计寿命",
    "功率密度": "功率密度", "质量": "质量", "尺寸": "外形尺寸",
    "电压范围": "电压范围", "电流范围": "电流范围", "氢气纯度": "氢气纯度",
    "电解槽功耗": "电解槽功耗", "负载范围": "负载范围", "热启动时间": "热启动时间", "冷启动时间": "冷启动时间",
}


def normalize_result_rows(results: list):
    rows = []
    for result in results or []:
        headers = result.get("table_headers") or []
        table_rows = result.get("table_rows") or []
        for row in table_rows:
            if not row: continue
            field = str(row[0])
            for idx, value in enumerate(row[1:], 1):
                rows.append({"standard": STANDARD_FIELDS.get(field, field), "raw": field, "model": headers[idx] if idx < len(headers) else "", "value": str(value), "doc": result.get("doc", ""), "page": result.get("page", ""), "status": "已引用原始表格"})
    return rows


def detect_conflicts(rows: list):
    grouped = {}
    for item in rows:
        key = (item.get("model") or "未标注型号", item.get("standard") or "未标注字段")
        grouped.setdefault(key, []).append(item)
    conflicts = []
    for (model, field), items in grouped.items():
        values = {str(x.get("value", "")).strip() for x in items if str(x.get("value", "")).strip()}
        if len(values) > 1:
            conflicts.append({"model": model, "field": field, "values": sorted(values), "sources": [{"doc":x.get("doc"), "page":x.get("page"), "value":x.get("value")} for x in items]})
    return conflicts


def proposal_payload(session_id: str, user: dict, version_id: str | None = None):
    with db().connect() as conn:
        params = {"s": session_id, "u": user["id"]}
        clause = "AND id=:v" if version_id else ""
        if version_id: params["v"] = version_id
        row = conn.execute(text(f"SELECT * FROM proposal_versions WHERE session_id=:s AND user_id=:u {clause} ORDER BY version_no DESC LIMIT 1"), params).mappings().first()
    if not row: raise HTTPException(404, "方案版本不存在")
    payload = row["payload"] if isinstance(row["payload"], dict) else json.loads(row["payload"])
    return dict(row), payload

class RegisterRequest(BaseModel):
    username: str
    email: str
    password: str
    user_type: str = "customer"
    admin_code: Optional[str] = None

class LoginRequest(BaseModel):
    username: str
    password: str

class SessionRequest(BaseModel):
    name: Optional[str] = None
    role: str = "customer_service"

class MessageRequest(BaseModel):
    role: str
    content: str
    metadata: Optional[dict] = None

def db():
    if not DATABASE_URL:
        raise HTTPException(503, "数据库未配置")
    return create_engine(DATABASE_URL)

def token_for(user):
    return jwt.encode({"sub": str(user["id"]), "username": user["username"], "exp": datetime.now(timezone.utc) + timedelta(days=7)}, SECRET_KEY, algorithm=ALGORITHM)

def current_user(authorization: Optional[str] = Header(None)):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(401, "请先登录")
    try:
        payload = jwt.decode(authorization[7:], SECRET_KEY, algorithms=[ALGORITHM])
        uid = payload.get("sub")
        if not uid: raise JWTError()
    except JWTError:
        raise HTTPException(401, "登录已失效，请重新登录")
    with db().connect() as conn:
        row = conn.execute(text("SELECT id, username, email, is_admin, user_type FROM users WHERE id=:id AND is_active=TRUE"), {"id": uid}).mappings().first()
    if not row: raise HTTPException(401, "用户不存在或已停用")
    return dict(row)

def ensure_chat_tables():
    with db().begin() as conn:
        conn.execute(text("""CREATE TABLE IF NOT EXISTS chat_sessions (id UUID PRIMARY KEY DEFAULT uuid_generate_v4(), user_id UUID REFERENCES users(id) ON DELETE CASCADE, session_name VARCHAR(255), assistant_role VARCHAR(32) DEFAULT 'customer_service', memory_summary TEXT, created_at TIMESTAMP DEFAULT NOW(), updated_at TIMESTAMP DEFAULT NOW())"""))
        conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS user_type VARCHAR(16) DEFAULT 'customer'"))
        conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS is_admin BOOLEAN DEFAULT FALSE"))
        conn.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS memory_summary TEXT"))
        conn.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS assistant_role VARCHAR(32) DEFAULT 'customer_service'"))
        conn.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS lead_score INTEGER DEFAULT 0"))
        conn.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS lead_level VARCHAR(16) DEFAULT 'low'"))
        conn.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS lead_signals JSONB DEFAULT '{}'::jsonb"))
        conn.execute(text("""CREATE TABLE IF NOT EXISTS chat_messages (id UUID PRIMARY KEY DEFAULT uuid_generate_v4(), session_id UUID REFERENCES chat_sessions(id) ON DELETE CASCADE, role VARCHAR(20) NOT NULL, content TEXT NOT NULL, metadata JSONB, created_at TIMESTAMP DEFAULT NOW())"""))
        conn.execute(text("""CREATE TABLE IF NOT EXISTS proposal_versions (id UUID PRIMARY KEY DEFAULT uuid_generate_v4(), session_id UUID REFERENCES chat_sessions(id) ON DELETE CASCADE, user_id UUID REFERENCES users(id) ON DELETE CASCADE, version_no INTEGER NOT NULL DEFAULT 1, title VARCHAR(255) NOT NULL, payload JSONB NOT NULL DEFAULT '{}'::jsonb, created_at TIMESTAMP DEFAULT NOW())"""))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_proposal_versions_session ON proposal_versions(session_id, created_at DESC)"))

@router.post("/sessions/{session_id}/proposals")
def save_proposal(session_id: str, req: ProposalSaveRequest, user=Depends(current_user)):
    normalized_rows = normalize_result_rows(req.results)
    conflicts = detect_conflicts(normalized_rows)
    for item in normalized_rows:
        if any(c["model"] == item["model"] and c["field"] == item["standard"] for c in conflicts):
            item["status"] = "存在参数冲突，待人工确认"
    payload = {"profile": req.profile, "content": req.content, "results": req.results, "normalized_rows": normalized_rows, "conflicts": conflicts}
    with db().begin() as conn:
        ok = conn.execute(text("SELECT 1 FROM chat_sessions WHERE id=:s AND user_id=:u"), {"s": session_id, "u": user["id"]}).first()
        if not ok: raise HTTPException(404, "会话不存在")
        no = conn.execute(text("SELECT COALESCE(MAX(version_no),0)+1 FROM proposal_versions WHERE session_id=:s"), {"s": session_id}).scalar()
        row = conn.execute(text("INSERT INTO proposal_versions(session_id,user_id,version_no,title,payload) VALUES(:s,:u,:n,:t,CAST(:p AS JSONB)) RETURNING id,version_no,title,created_at"), {"s":session_id,"u":user["id"],"n":no,"t":req.title,"p":json.dumps(payload,ensure_ascii=False)}).mappings().first()
    return {"proposal": dict(row), "normalized_count": len(payload["normalized_rows"]), "conflicts": conflicts}

@router.get("/sessions/{session_id}/proposals/{version_id}")
def get_proposal(session_id: str, version_id: str, user=Depends(current_user)):
    row, payload = proposal_payload(session_id, user, version_id)
    return {"proposal": {"id": row["id"], "version_no": row["version_no"], "title": row["title"], "created_at": row["created_at"], "payload": payload}}

@router.get("/sessions/{session_id}/proposals/{version_id}/compare/{other_version_id}")
def compare_proposals(session_id: str, version_id: str, other_version_id: str, user=Depends(current_user)):
    left, a = proposal_payload(session_id, user, version_id); right, b = proposal_payload(session_id, user, other_version_id)
    keys = sorted(set((a.get("profile") or {}) | (b.get("profile") or {})))
    profile_changes = [{"field": k, "from": (b.get("profile") or {}).get(k), "to": (a.get("profile") or {}).get(k)} for k in keys if (a.get("profile") or {}).get(k) != (b.get("profile") or {}).get(k)]
    return {"base_version": right["version_no"], "target_version": left["version_no"], "profile_changes": profile_changes, "content_changed": (a.get("content") or "") != (b.get("content") or ""), "conflict_count": len(a.get("conflicts") or [])}

@router.get("/sessions/{session_id}/proposals")
def list_proposals(session_id: str, user=Depends(current_user)):
    with db().connect() as conn:
        rows = conn.execute(text("SELECT id,version_no,title,created_at FROM proposal_versions WHERE session_id=:s AND user_id=:u ORDER BY version_no DESC"), {"s":session_id,"u":user["id"]}).mappings().all()
    return {"proposals": [dict(x) for x in rows]}

@router.get("/sessions/{session_id}/proposals/{version_id}/docx")
def proposal_docx(session_id: str, version_id: str, user=Depends(current_user)):
    row, p = proposal_payload(session_id, user, version_id)
    doc = Document(); doc.add_heading(p.get("profile", {}).get("项目名称", row["title"]), 0)
    doc.add_paragraph(f"方案版本：V{row['version_no']}    生成时间：{row['created_at']}")
    doc.add_heading("项目需求画像", 1)
    for k,v in (p.get("profile") or {}).items(): doc.add_paragraph(f"{k}：{v}")
    doc.add_heading("技术方案初稿", 1); doc.add_paragraph(p.get("content") or "暂无方案正文")
    doc.add_heading("参数与来源", 1)
    table = doc.add_table(rows=1, cols=6); table.style = "Table Grid"
    for cell, val in zip(table.rows[0].cells, ["标准字段","原始字段","型号","原始值","来源","核验状态"]): cell.text=val
    for x in p.get("normalized_rows", []):
        cells=table.add_row().cells
        for c,val in zip(cells,[x.get("standard"),x.get("raw"),x.get("model"),x.get("value"),f"{x.get('doc')} P{x.get('page')}",x.get("status")]): c.text=str(val or "")
    out=BytesIO(); doc.save(out); out.seek(0)
    return StreamingResponse(out, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=proposal_v{row['version_no']}.docx"})

@router.get("/sessions/{session_id}/proposals/{version_id}/xlsx")
def proposal_xlsx(session_id: str, version_id: str, user=Depends(current_user)):
    row, p = proposal_payload(session_id, user, version_id)
    headers = ["标准字段","原始字段","型号","原始值","来源文档","页码","核验状态"]
    rows = [headers] + [[x.get("standard"),x.get("raw"),x.get("model"),x.get("value"),x.get("doc"),x.get("page"),x.get("status")] for x in p.get("normalized_rows", [])]
    def cell(value): return f'<c t="inlineStr"><is><t>{html.escape(str(value or ""))}</t></is></c>'
    sheet_rows = ''.join(f'<row r="{i}">' + ''.join(cell(v) for v in values) + '</row>' for i, values in enumerate(rows, 1))
    xml = f'<?xml version="1.0" encoding="UTF-8"?><worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"><sheetData>{sheet_rows}</sheetData></worksheet>'
    content = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/><Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/></Types>'
    rels = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/></Relationships>'
    workbook = '<?xml version="1.0" encoding="UTF-8"?><workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><sheets><sheet name="参数对比" sheetId="1" r:id="rId1"/></sheets></workbook>'
    wbrels = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/></Relationships>'
    out=BytesIO()
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as z:
        z.writestr('[Content_Types].xml',content); z.writestr('_rels/.rels',rels); z.writestr('xl/workbook.xml',workbook); z.writestr('xl/_rels/workbook.xml.rels',wbrels); z.writestr('xl/worksheets/sheet1.xml',xml)
    out.seek(0)
    return StreamingResponse(out, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": f"attachment; filename=proposal_v{row['version_no']}.xlsx"})

@router.get("/sessions/{session_id}/proposals/{version_id}/pdf")
def proposal_pdf(session_id: str, version_id: str, user=Depends(current_user)):
    row, p = proposal_payload(session_id, user, version_id); pdf = fitz.open(); page = pdf.new_page(width=595, height=842)
    y=50; page.insert_text((50,y), row["title"], fontsize=16, fontname="china-s"); y+=28
    page.insert_text((50,y), f"方案版本：V{row['version_no']}", fontsize=10, fontname="china-s"); y+=24
    sections=[("项目需求画像", "；".join(f"{k}：{v}" for k,v in (p.get("profile") or {}).items())), ("技术方案初稿",p.get("content") or "暂无方案正文")]
    for heading, body in sections:
        if y>770: page=pdf.new_page(width=595,height=842); y=50
        page.insert_text((50,y),heading,fontsize=13,fontname="china-s"); y+=22
        for raw_line in str(body).splitlines() or [""]:
            for line in [raw_line[i:i+42] for i in range(0,len(raw_line),42)] or [""]:
                if y>790: page=pdf.new_page(width=595,height=842); y=50
                page.insert_text((55,y),line,fontsize=9,fontname="china-s"); y+=14
        y+=10
    out=BytesIO(); pdf.save(out); pdf.close(); out.seek(0)
    return StreamingResponse(out, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=proposal_v{row['version_no']}.pdf"})


@router.post("/register")
def register(req: RegisterRequest):
    if len(req.username.strip()) < 2 or len(req.password) < 6:
        raise HTTPException(400, "用户名至少2位，密码至少6位")
    ensure_chat_tables()
    user_type = req.user_type if req.user_type in ("customer", "admin") else "customer"
    if user_type == "admin":
        expected = os.getenv("ADMIN_REGISTRATION_CODE", "qingpu-admin-2026")
        if not req.admin_code or req.admin_code != expected:
            raise HTTPException(403, "管理用户注册口令错误")
    try:
        with db().begin() as conn:
            row = conn.execute(text("INSERT INTO users(username,email,hashed_password,user_type,is_admin) VALUES(:u,:e,:p,:t,:a) RETURNING id,username,email,is_admin,user_type"), {"u":req.username.strip(),"e":req.email.strip(),"p":pwd_context.hash(req.password),"t":user_type,"a":user_type == "admin"}).mappings().first()
        user=dict(row)
        return {"token": token_for(user), "user": user}
    except Exception as e:
        if "unique" in str(e).lower(): raise HTTPException(409, "用户名或邮箱已存在")
        raise HTTPException(500, "注册失败")

@router.post("/login")
def login(req: LoginRequest):
    with db().connect() as conn:
        row=conn.execute(text("SELECT id,username,email,hashed_password,is_admin,user_type FROM users WHERE username=:u OR email=:u"), {"u":req.username.strip()}).mappings().first()
    if not row or not pwd_context.verify(req.password, row["hashed_password"]): raise HTTPException(401,"用户名或密码错误")
    user={k:row[k] for k in ("id","username","email","is_admin","user_type")}
    return {"token":token_for(user),"user":user}

@router.get("/me")
def me(user=Depends(current_user)): return {"user":user}

@router.get("/sessions")
def sessions(user=Depends(current_user)):
    with db().connect() as conn:
        rows=conn.execute(text("SELECT id,session_name,assistant_role,lead_score,lead_level,lead_signals,created_at,updated_at FROM chat_sessions WHERE user_id=:u ORDER BY updated_at DESC"),{"u":user["id"]}).mappings().all()
    return {"sessions":[dict(r) for r in rows]}

@router.get("/sessions/{session_id}")
def get_session(session_id: str, user=Depends(current_user)):
    """加载单个会话及完整消息历史，供前端历史记录抽屉恢复上下文。"""
    with db().connect() as conn:
        session = conn.execute(text("SELECT id,session_name,assistant_role,created_at,updated_at FROM chat_sessions WHERE id=:s AND user_id=:u"), {"s":session_id,"u":user["id"]}).mappings().first()
        if not session: raise HTTPException(404, "会话不存在")
        rows = conn.execute(text("SELECT id,role,content,metadata,created_at FROM chat_messages WHERE session_id=:s ORDER BY created_at"), {"s":session_id}).mappings().all()
    data = dict(session); data["messages"] = [dict(row) for row in rows]
    return data

@router.post("/sessions")
def create_session(req: SessionRequest, user=Depends(current_user)):
    with db().begin() as conn:
        row=conn.execute(text("INSERT INTO chat_sessions(user_id,session_name,assistant_role) VALUES(:u,:n,:r) RETURNING id,session_name,assistant_role,created_at,updated_at"),{"u":user["id"],"n":req.name or "新会话","r":req.role if req.role in ("customer_service","sales","technical_support") else "customer_service"}).mappings().first()
    return {"session":dict(row)}

@router.put("/sessions/{session_id}/role")
def update_session_role(session_id: str, req: SessionRequest, user=Depends(current_user)):
    role = req.role if req.role in ("customer_service", "sales", "technical_support") else "customer_service"
    with db().begin() as conn:
        row = conn.execute(text("UPDATE chat_sessions SET assistant_role=:r, updated_at=NOW() WHERE id=:s AND user_id=:u RETURNING assistant_role"), {"r":role,"s":session_id,"u":user["id"]}).mappings().first()
        if not row: raise HTTPException(404, "会话不存在")
    return {"assistant_role": row["assistant_role"]}

@router.get("/sessions/{session_id}/lead")
def session_lead(session_id: str, user=Depends(current_user)):
    with db().connect() as conn:
        row=conn.execute(text("SELECT lead_score,lead_level,lead_signals FROM chat_sessions WHERE id=:s AND user_id=:u"), {"s":session_id,"u":user["id"]}).mappings().first()
    if not row: raise HTTPException(404, "会话不存在")
    return dict(row)

@router.get("/sessions/{session_id}/messages")
def get_messages(session_id: str, user=Depends(current_user)):
    with db().connect() as conn:
        ok=conn.execute(text("SELECT 1 FROM chat_sessions WHERE id=:s AND user_id=:u"),{"s":session_id,"u":user["id"]}).first()
        if not ok: raise HTTPException(404,"会话不存在")
        rows=conn.execute(text("SELECT id,role,content,metadata,created_at FROM chat_messages WHERE session_id=:s ORDER BY created_at"),{"s":session_id}).mappings().all()
    return {"messages":[dict(r) for r in rows]}

@router.post("/sessions/{session_id}/messages")
def save_message(session_id: str, req: MessageRequest, user=Depends(current_user)):
    with db().begin() as conn:
        ok=conn.execute(text("SELECT 1 FROM chat_sessions WHERE id=:s AND user_id=:u"),{"s":session_id,"u":user["id"]}).first()
        if not ok: raise HTTPException(404,"会话不存在")
        row=conn.execute(text("INSERT INTO chat_messages(session_id,role,content,metadata) VALUES(:s,:r,:c,CAST(:m AS JSONB)) RETURNING id,created_at"),{"s":session_id,"r":req.role,"c":req.content,"m":json.dumps(req.metadata or {}, ensure_ascii=False)}).mappings().first()
        conn.execute(text("UPDATE chat_sessions SET updated_at=NOW() WHERE id=:s"),{"s":session_id})
    return {"message":dict(row)}

@router.get("/sessions/{session_id}/memory")
def session_memory(session_id: str, user=Depends(current_user)):
    """返回长期摘要与最近 6 条原文，供下一轮模型调用。"""
    with db().connect() as conn:
        session = conn.execute(text("SELECT memory_summary FROM chat_sessions WHERE id=:s AND user_id=:u"), {"s":session_id,"u":user["id"]}).mappings().first()
        if not session: raise HTTPException(404, "会话不存在")
        rows = conn.execute(text("SELECT role,content FROM chat_messages WHERE session_id=:s ORDER BY created_at DESC LIMIT 6"), {"s":session_id}).mappings().all()
    history = []
    if session["memory_summary"]: history.append({"role":"system","content":"以下是此前对话的长期记忆，请据此保持上下文一致：\n" + session["memory_summary"]})
    history += [{"role":r["role"],"content":r["content"]} for r in reversed(rows)]
    return {"history": history}

@router.post("/sessions/{session_id}/compact")
async def compact_session(session_id: str, user=Depends(current_user)):
    """消息超过阈值时压缩旧上下文；原始消息保留用于用户查看。"""
    from app.services.llm_service import llm_service
    with db().connect() as conn:
        session = conn.execute(text("SELECT memory_summary FROM chat_sessions WHERE id=:s AND user_id=:u"), {"s":session_id,"u":user["id"]}).mappings().first()
        if not session: raise HTTPException(404, "会话不存在")
        rows = conn.execute(text("SELECT role,content FROM chat_messages WHERE session_id=:s ORDER BY created_at"), {"s":session_id}).mappings().all()
    if len(rows) <= 12: return {"compacted":False}
    old = rows[:-6]
    transcript = "\n".join(f"{'用户' if r['role']=='user' else '助手'}：{r['content']}" for r in old)
    summary = await llm_service.summarize_memory(session["memory_summary"] or "", transcript)
    with db().begin() as conn:
        conn.execute(text("UPDATE chat_sessions SET memory_summary=:m, updated_at=NOW() WHERE id=:s"), {"m":summary,"s":session_id})
    return {"compacted":True}

@router.delete("/sessions/{session_id}")
def delete_session(session_id: str, user=Depends(current_user)):
    with db().begin() as conn:
        result=conn.execute(text("DELETE FROM chat_sessions WHERE id=:s AND user_id=:u"),{"s":session_id,"u":user["id"]})
    if result.rowcount==0: raise HTTPException(404,"会话不存在")
    return {"ok":True}
